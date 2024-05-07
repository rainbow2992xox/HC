from docx import Document


def get_last_level_headings_and_text(doc):
    headings_and_text = []
    current_heading = None
    current_text = []

    for paragraph in doc.paragraphs:
        # Check if the paragraph is a heading
        if paragraph.style.name.startswith('Heading'):
            # Extract the level of the heading
            heading_level = int(paragraph.style.name.split()[1])  # Extracting the number from style name like "Heading 1"

            # Check if it's the last level heading
            if heading_level == 1:
                # If there is previous heading and text, add it to the list
                if current_heading is not None:
                    headings_and_text.append((current_heading, '\n'.join(current_text)))
                    current_text = []

                # Update current heading
                current_heading = paragraph.text
            else:
                # For sub-level headings, append their text to current_text
                current_text.append(paragraph.text)
        else:
            # If not a heading, add text to current_text list
            current_text.append(paragraph.text)

    # Add the last heading and text
    if current_heading is not None:
        headings_and_text.append((current_heading, '\n'.join(current_text)))

    return headings_and_text


# Load the document
doc = Document('test.docx')  # Replace 'your_document.docx' with the path to your document
headings_and_text = get_last_level_headings_and_text(doc)

# Print headings and text
for heading, text in headings_and_text:
    print("Heading:", heading)
    print("Text:", text)
    print("----------")


tables = doc.tables

# 遍历每个表格
for table in tables:
    # 遍历表格的每一行
    for row in table.rows:
        # 遍历行中的每个单元格
        for cell in row.cells:
            # 打印单元格中的文本
            print(cell.text)
